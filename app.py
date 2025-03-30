import os
import base64
import logging
from logging.handlers import RotatingFileHandler
from datetime import datetime
import math
from io import BytesIO
import mimetypes
import filetype

from docx import Document as DocxDocument
from pdf2image import convert_from_bytes
from flask import (
    Flask, render_template, request, redirect,
    url_for, send_from_directory, flash, abort, jsonify
)
from flask_login import (
    LoginManager, login_user,
    login_required, logout_user, current_user
)
from flask_bcrypt import Bcrypt
from flask_caching import Cache
from flask_wtf.csrf import CSRFProtect
from flask_migrate import Migrate
from werkzeug.utils import secure_filename
from dotenv import load_dotenv

# Ortam değişkenlerini .env dosyasından yükleyelim.
load_dotenv()

# Database instance'ını extensions modülünden alıyoruz.
from extensions import db

def configure_logging():
    log_dir = os.path.join(os.path.dirname(__file__), 'logs')
    os.makedirs(log_dir, exist_ok=True)
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        handlers=[
            RotatingFileHandler(os.path.join(log_dir, 'app.log'), encoding='utf-8', maxBytes=10_000_000, backupCount=5),
            logging.StreamHandler()
        ]
    )
    if os.name == 'nt':
        try:
            from ctypes import windll
            windll.kernel32.SetConsoleOutputCP(65001)
        except Exception:
            pass

configure_logging()
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.config.from_object('config.Config')
app.config['SQLALCHEMY_DATABASE_URI'] = f"sqlite:///{os.path.join(os.path.dirname(__file__), 'instance', 'app.db')}"
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Initialize extensions
db.init_app(app)
migrate = Migrate(app, db)
csrf = CSRFProtect(app)
bcrypt = Bcrypt(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'
cache = Cache(app, config={'CACHE_TYPE': 'SimpleCache'})

# Modelleri içe aktaralım.
from models import User, Document, DocumentSequence

@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))

def admin_required(f):
    from functools import wraps
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or current_user.role != 'admin':
            abort(403)
        return f(*args, **kwargs)
    return decorated_function

def guess_mime_from_bytes(data: bytes):
    kind = filetype.guess(data)
    return kind.mime if kind else None

def validate_file(file_storage):
    filename = file_storage.filename
    ext = os.path.splitext(filename)[1].lower().strip('.')
    if ext not in app.config['ALLOWED_EXTENSIONS']:
        raise ValueError(f"Disallowed file extension: {ext}")
    file_storage.seek(0)
    data = file_storage.read()
    file_storage.seek(0)
    detected_mime = guess_mime_from_bytes(data)
    if not detected_mime:
        raise ValueError("Could not guess file type from content.")
    if detected_mime not in app.config['ALLOWED_MIME_TYPES']:
        raise ValueError(f"Disallowed file type: {detected_mime}")
    return True

def log_activity(action, filename=None):
    logger.info("User: %s - Action: %s - File: %s", 
                current_user.username if current_user.is_authenticated else "Anonymous", 
                action, filename or 'None')

@app.route('/')
def index():
    if current_user.is_authenticated:
        return redirect(url_for('documents'))
    return redirect(url_for('login'))

@app.route('/documents')
@login_required
def documents():
    search_query = request.args.get('query', '').lower()
    filter_category = request.args.get('filter_category', '').strip()
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 10, type=int)
    
    query = Document.query
    if search_query:
        query = query.filter(
            db.or_(
                Document.filename.ilike(f'%{search_query}%'),
                Document.description.ilike(f'%{search_query}%')
            )
        )
    if filter_category:
        query = query.filter_by(category=filter_category)
    
    total_docs = query.count()
    documents = query.order_by(Document.upload_date.desc()) \
                     .offset((page - 1) * per_page) \
                     .limit(per_page) \
                     .all()
    total_pages = math.ceil(total_docs / per_page) if total_docs else 1
    return render_template('index.html',
                           documents=documents,
                           search_query=search_query,
                           filter_category=filter_category,
                           current_page=page,
                           total_pages=total_pages,
                           per_page=per_page)

@app.route('/preview/<filename>')
@login_required
@cache.memoize(timeout=3600)
def preview_file(filename):
    safe_filename = secure_filename(filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], safe_filename)
    if not os.path.exists(file_path):
        flash('File not found', 'error')
        return redirect(url_for('documents'))
    
    mime_type = None
    try:
        kind = filetype.guess(file_path)
        if kind:
            mime_type = kind.mime
        if not mime_type:
            guess, _ = mimetypes.guess_type(file_path)
            mime_type = guess or 'application/octet-stream'
    except Exception as e:
        logger.error("Error detecting file type in preview: %s", str(e))
        mime_type = 'application/octet-stream'
    
    if mime_type == 'application/pdf':
        try:
            with open(file_path, 'rb') as f:
                pdf_bytes = f.read()
            images = convert_from_bytes(
                pdf_bytes,
                first_page=1,
                last_page=1,
                dpi=100,
                poppler_path=app.config['POPPLER_PATH']
            )
            buffered = BytesIO()
            images[0].save(buffered, format="JPEG")
            img_str = base64.b64encode(buffered.getvalue()).decode()
            return render_template('preview.html', content=f'data:image/jpeg;base64,{img_str}', filename=safe_filename)
        except Exception as e:
            logger.error("PDF preview error: %s", str(e))
            flash('PDF preview error', 'error')
            return redirect(url_for('documents'))
    elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        try:
            doc = DocxDocument(file_path)
            full_text = [para.text for para in doc.paragraphs if para.text.strip()]
            preview_content = '<br>'.join(full_text[:20])
            return render_template('preview.html', content=preview_content, filename=safe_filename)
        except Exception as e:
            logger.error("DOCX preview error: %s", str(e))
            flash('DOCX preview error', 'error')
            return redirect(url_for('documents'))
    elif mime_type in ['image/jpeg', 'image/png']:
        try:
            with open(file_path, 'rb') as f:
                image_data = f.read()
            img_str = base64.b64encode(image_data).decode()
            return render_template('preview.html', content=f'data:{mime_type};base64,{img_str}', filename=safe_filename)
        except Exception as e:
            logger.error("Image preview error: %s", str(e))
            flash('Image preview error', 'error')
            return redirect(url_for('documents'))
    elif mime_type.startswith('application/') or mime_type.startswith('image/'):
        flash('Bu dosya türü için önizleme desteklenmiyor.', 'info')
        return redirect(url_for('documents'))
    else:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read(2000)
            return render_template('preview.html', content=content, filename=safe_filename)
        except Exception as e:
            logger.error("Text preview error: %s", str(e))
            flash('Text preview error', 'error')
            return redirect(url_for('documents'))

@app.route('/admin/upload', methods=['GET'])
@login_required
@admin_required
def admin_upload_form():
    return render_template('upload.html')

@app.route('/admin/users')
@login_required
@admin_required
def admin_users():
    users = User.query.order_by(User.id).all()
    return render_template('users.html', users=users)

@app.route('/create-user', methods=['POST'])
@login_required
@admin_required
def create_user_route():
    if not all(k in request.form for k in ['username', 'password', 'role']):
        abort(400)
    
    hashed_pw = bcrypt.generate_password_hash(request.form['password']).decode('utf-8')
    new_user = User(
        username=request.form['username'],
        password=hashed_pw,
        role=request.form['role'],
        first_name=request.form.get('first_name'),
        last_name=request.form.get('last_name')
    )
    try:
        db.session.add(new_user)
        db.session.commit()
        flash('Kullanıcı başarıyla oluşturuldu', 'success')
    except Exception as e:
        db.session.rollback()
        flash('Bu kullanıcı adı zaten alınmış veya başka bir hata oluştu', 'error')
    return redirect(url_for('admin_users'))

@app.route('/admin/users/change_role', methods=['POST'])
@login_required
@admin_required
def change_role():
    user_id = request.form.get('user_id', '').strip()
    new_role = request.form.get('role', 'user').strip()
    if not user_id:
        flash('Kullanıcı ID eksik', 'error')
        return redirect(url_for('admin_users'))
    try:
        user = User.query.get(int(user_id))
        if user:
            user.role = new_role
            db.session.commit()
            flash('Kullanıcı rolü güncellendi', 'success')
        else:
            flash('Kullanıcı bulunamadı', 'error')
    except Exception as e:
        db.session.rollback()
        flash('Kullanıcı rolü güncelleme hatası: ' + str(e), 'error')
    return redirect(url_for('admin_users'))

@app.route('/admin/users/delete', methods=['POST'])
@login_required
@admin_required
def delete_user():
    user_id = request.form.get('user_id', '').strip()
    if not user_id:
        flash('Kullanıcı ID bulunamadı', 'error')
        return redirect(url_for('admin_users'))
    if user_id == str(current_user.id):
        flash('Kendi hesabınızı silemezsiniz!', 'error')
        return redirect(url_for('admin_users'))
    try:
        user = User.query.get(int(user_id))
        if user:
            db.session.delete(user)
            db.session.commit()
            flash('Kullanıcı silindi', 'success')
        else:
            flash('Kullanıcı bulunamadı', 'error')
    except Exception as e:
        db.session.rollback()
        flash('Kullanıcı silme hatası: ' + str(e), 'error')
    return redirect(url_for('admin_users'))

@app.route('/upload', methods=['POST'])
@login_required
@admin_required
def upload_file():
    if 'file' not in request.files:
        flash('No file selected', 'error')
        return redirect(url_for('documents'))
    file = request.files['file']
    if file.filename == '':
        flash('No file selected', 'error')
        return redirect(url_for('documents'))
    try:
        validate_file(file)
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        counter = 1
        while os.path.exists(file_path):
            base, ext = os.path.splitext(filename)
            filename = f"{base}_{counter}{ext}"
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            counter += 1
        file.save(file_path)

        originator = request.form.get('originator')
        document_type = request.form.get('document_type')
        discipline = request.form.get('discipline')
        building_code = request.form.get('building_code')
        category = request.form.get('category', '')
        revision = 0

        doc_number = generate_document_number(originator, document_type, discipline, building_code, revision)

        new_doc = Document(
            filename=filename,
            description=request.form.get('description', ''),
            upload_date=datetime.utcnow(),
            user_id=current_user.id,
            category=category,
            project_code="SPP2",
            revision=revision,
            document_number=doc_number
        )
        db.session.add(new_doc)
        db.session.commit()
        log_activity('upload', filename)
        flash('File uploaded successfully', 'success')
    except ValueError as ve:
        flash(str(ve), 'error')
    except Exception as e:
        db.session.rollback()
        logger.error("Upload error: %s", str(e))
        flash('File upload error', 'error')
    return redirect(url_for('documents'))

def generate_document_number(originator, document_type, discipline, building_code, revision):
    # Get the most recent policy for numbering; since DocumentPolicy is removed,
    # use default values for project_code and revision_format.
    default_project_code = "SPP2"
    default_revision_format = "R%02d"
    
    # Get or create DocumentSequence record for the combination.
    seq = DocumentSequence.query.filter_by(
        originator=originator,
        document_type=document_type,
        discipline=discipline,
        building_code=building_code
    ).first()
    if not seq:
        seq = DocumentSequence(
            originator=originator,
            document_type=document_type,
            discipline=discipline,
            building_code=building_code,
            next_sequence=1
        )
        db.session.add(seq)
        db.session.flush()
    sequence_number = seq.next_sequence
    seq.next_sequence += 1

    revision_str = default_revision_format % revision
    document_number = f"{originator}-{default_project_code}-{document_type}-{discipline}-{building_code}-{sequence_number:03d}_{revision_str}"
    return document_number

@app.route('/download/<filename>')
@login_required
def download_file(filename):
    safe_filename = secure_filename(filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], safe_filename)
    if not os.path.exists(file_path):
        flash('File not found', 'error')
        return redirect(url_for('documents'))
    doc = Document.query.filter_by(filename=safe_filename).first()
    if not doc:
        abort(403)
    log_activity('download', safe_filename)
    return send_from_directory(app.config['UPLOAD_FOLDER'], safe_filename, as_attachment=True)

@app.route('/delete/<filename>')
@login_required
@admin_required
def delete_file(filename):
    safe_filename = secure_filename(filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], safe_filename)
    doc = Document.query.filter_by(filename=safe_filename).first()
    if not doc:
        flash('File not found in database', 'error')
        return redirect(url_for('documents'))
    try:
        db.session.delete(doc)
        db.session.commit()
        if os.path.exists(file_path):
            os.remove(file_path)
        log_activity('delete', safe_filename)
        flash('File deleted successfully', 'success')
    except Exception as e:
        db.session.rollback()
        flash('File deletion error', 'error')
    return redirect(url_for('documents'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('documents'))
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = User.query.filter_by(username=username).first()
        if user and bcrypt.check_password_hash(user.password, password):
            login_user(user)
            log_activity('login')
            return redirect(url_for('documents'))
        flash('Invalid username or password', 'error')
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    log_activity('logout')
    logout_user()
    return redirect(url_for('login'))

@app.route('/favicon.ico')
def favicon():
    return send_from_directory(os.path.join(app.root_path, 'static'),
                               'favicon.ico', mimetype='image/vnd.microsoft.icon')

@app.before_request
def before_request():
    if (not current_user.is_authenticated and
        request.endpoint not in ['login', 'static'] and
        not request.path.startswith('/static/')):
        return redirect(url_for('login', next=request.url))

@app.errorhandler(403)
def forbidden(e):
    return render_template('403.html'), 403

@app.errorhandler(404)
def page_not_found(e):
    if request.accept_mimetypes.accept_json and not request.accept_mimetypes.accept_html:
        return jsonify({'error': 'Not found'}), 404
    return render_template('404.html'), 404

@app.errorhandler(500)
def internal_server_error(e):
    logger.error(f"Server error: {str(e)}")
    if request.accept_mimetypes.accept_json and not request.accept_mimetypes.accept_html:
        return jsonify({'error': 'Internal server error'}), 500
    return render_template('500.html'), 500

if __name__ == '__main__':
    os.makedirs(app.config.get('UPLOAD_FOLDER', 'uploads'), exist_ok=True)
    app.run(host='0.0.0.0', port=5000, debug=app.config['DEBUG'])
